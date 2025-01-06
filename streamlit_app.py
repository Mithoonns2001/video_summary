import subprocess
import whisper
import google.generativeai as genai
import streamlit as st
import tempfile
import os
from moviepy.editor import VideoFileClip, TextClip, CompositeVideoClip
import json
import cv2
from docx import Document
from docx.shared import Inches
import shutil
# Get the current working directory
current_directory = os.getcwd()

# Function to extract audio from video using FFmpeg
def extract_audio_from_video(video_path, audio_path):
    """
    Extract audio from a video file using FFmpeg.
    """
    command = [
        "ffmpeg", "-i", video_path, "-vn", "-acodec", "pcm_s16le",
        "-ar", "16000", "-ac", "1", audio_path
    ]
    subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    st.write(f"Audio extracted to {audio_path}")

# Function to transcribe audio using Whisper
def transcribe_audio(audio_path):
    model = whisper.load_model("tiny")  # Options: tiny, base, small, medium, large
    result = model.transcribe(audio_path)
    
    transcribed_text = ""
    for segment in result["segments"]:
        start = segment["start"]
        end = segment["end"]
        text = segment["text"]
        transcribed_text += f"[{start:.2f} - {end:.2f}] {text}\n"

    return transcribed_text

# Function to generate summary using Gemini model
def generate_summary(transcribed_text):
    genai.configure(api_key="AIzaSyAqfhjwbo1tLpdGwsEI9DNW6Jp6NcpFIY8")
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = f"""
    provide combined summary for each important part with timestamp in format [{{
        time:start time(in float format not in string and end time is not required), content: 'summary'
    }} for the following content:

    """ + transcribed_text

    response = model.generate_content(prompt)
    return response.text

def merge_summary_with_video(video_path, summary_json, output_video_path):
    """
    Overlay summary text on the video based on timestamps.
    """
    # Load the video
    video = VideoFileClip(video_path)

    # Parse summary JSON
    try:
        summary_data = json.loads(summary_json)
    except json.JSONDecodeError:
        st.error("Invalid JSON format in summary.")
        return

    max_font_size = video.h * 100  # Adjust this for desired size

    # Create a list of text clips
    text_clips = []
    for item in summary_data:
        start_time = item['start_time']
        end_time = item['end_time']
        content = item['content']

        text_clip = (
            TextClip(content, fontsize=max_font_size, color='white', bg_color='black', size=(video.w*5, video.h))
            .set_position(("center", "bottom")) #.set_position(("center", "bottom"))
            .set_start(start_time)
            .set_duration(end_time-start_time)  # Display each text for 5 seconds
        )

        text_clips.append(text_clip)

    # Combine the text clips with the video
    final_video = CompositeVideoClip([video, *text_clips])

    # Write the final video to the output path
    final_video.write_videofile(output_video_path, codec="libx264", audio_codec="aac")

    st.success(f"Summary overlaid on video and saved to {output_video_path}")


# Function to extract keyframes based on scene changes
def extract_keyframes(video_path, threshold=0.1):
    cap = cv2.VideoCapture(video_path)
    prev_hist = None
    prev_second = 0
    keyframe_seconds = []

    while True:
        ret, frame = cap.read()
        if not ret:
            break

        current_second = cap.get(cv2.CAP_PROP_POS_FRAMES) // cap.get(cv2.CAP_PROP_FPS)
        if current_second != prev_second:
            hist = cv2.calcHist([frame], [0, 1, 2], None, [8, 8, 8], [0, 256, 0, 256, 0, 256])
            hist = cv2.normalize(hist, hist).flatten()
            if prev_hist is not None:
                diff = cv2.compareHist(prev_hist, hist, cv2.HISTCMP_BHATTACHARYYA)
                if diff > threshold:
                    keyframe_seconds.append(int(prev_second))
            prev_hist = hist
            prev_second = current_second

    max_seconds = int(cap.get(cv2.CAP_PROP_FRAME_COUNT)) / cap.get(cv2.CAP_PROP_FPS)
    keyframe_seconds.append(int(max_seconds))
    cap.release()
    return keyframe_seconds

# Function to save keyframes as images
def save_keyframes(video_path, keyframe_seconds, output_folder):
    if os.path.exists(output_folder):
        shutil.rmtree(output_folder)
    os.makedirs(output_folder)

    cap = cv2.VideoCapture(video_path)
    for sec in keyframe_seconds:
        cap.set(cv2.CAP_PROP_POS_FRAMES, sec * cap.get(cv2.CAP_PROP_FPS))
        ret, frame = cap.read()
        if ret:
            cv2.imwrite(os.path.join(output_folder, f"{sec}.jpg"), frame)
    cap.release()
    return output_folder

# Function to generate a Word document with content and images
def generate_document(content_data, image_folder, output_path):
    doc = Document()
    image_names = sorted(os.listdir(image_folder), key=lambda x: int(x.split('.')[0]))
    image_num = [int(num.split('.')[0]) for num in image_names]
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("VIDEO SUMMARY")
    title_run.bold = True
    title_paragraph.alignment = 1 # Centwr Align

    for entry in content_data:
        start_time = entry['start_time']
        end_time = entry['end_time']
        content = entry['content']

        doc.add_heading(f"Summary Between ({start_time:.1f}s - {end_time:.1f}s):", level=2)
        image_paragraph = doc.add_paragraph()
        doc.add_paragraph(content)

        associated_images = [num for num in image_num if start_time <= num <= end_time]
        if associated_images:
            for img in associated_images:
                image_path = os.path.join(image_folder, f"{img}.jpg")
                image_paragraph.add_run().add_picture(image_path, width=Inches(3))

    doc.save(output_path)

# Streamlit UI setup
st.title("Video Transcription and Summarization App")
st.write("Upload a video to extract audio, transcribe it, and get a summarized content with timestamps.")

# Upload video
uploaded_video = st.file_uploader("Upload a Video File", type=["mp4", "mov", "avi"])

if uploaded_video is not None:
    # Save the uploaded video to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_video_file:
        temp_video_file.write(uploaded_video.read())
        temp_video_file_path = temp_video_file.name

    # Display the uploaded video
    # st.write(temp_video_file_path)  # This will display the video in the app
    st.video(temp_video_file_path)  # This will display the video in the app

    # Extract audio from the video
    audio_path = os.path.join(current_directory, "audio.wav")
    extract_audio_from_video(temp_video_file_path, audio_path)

    # Transcribe the extracted audio
    transcribed_text = transcribe_audio(audio_path)
    # st.text_area("Transcribed Text", transcribed_text, height=300)

    os.remove(audio_path)
    # Generate summary from the transcribed text
    # if st.button("Generate Summary"):
    summary = generate_summary(transcribed_text)
    # st.text_area("Generated Summary", summary, height=300)
    # Parse the summary JSON
    import re
    import json
    json_match = re.search(r'\[.*\]', summary, re.DOTALL)
    if json_match:
        json_data = json_match.group(0)  # Extracts the matched JSON string
        summary_data = json.loads(json_data)  # Converts the string into a Python list/dict
    else:
        st.write("No JSON found.")

    new_content_data=[]
    duration= 1010.9

    # Get the name of the video file
    # video_name = os.path.basename(uploaded_video)

    # Open the video using cv2
    cap = cv2.VideoCapture(temp_video_file_path)

    # Get the frame rate (frames per second)
    fps = cap.get(cv2.CAP_PROP_FPS)

    # Get the total number of frames
    frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT)

    # Calculate the duration in seconds
    duration_seconds = frame_count / fps if fps else 0
    
    for i in range(len(summary_data)-1):
        new_content_data.append({'start_time':summary_data[i]['time'], 'end_time':summary_data[i+1]['time'], 'content': summary_data[i]['content']})
    new_content_data.append({'start_time':summary_data[-1]['time'], 'end_time':duration_seconds, 'content': summary_data[-1]['content']})
    st.text_area("Generated Summary", new_content_data)

    if st.button("Generate Document"):

        keyframe_seconds = extract_keyframes(temp_video_file_path)
        image_folder = os.path.join(current_directory, "frames")
        save_keyframes(temp_video_file_path, keyframe_seconds, image_folder)

        output_document_path = os.path.join(current_directory, "Content_with_Images.docx")
        generate_document(new_content_data, image_folder, output_document_path)

        st.success(f"Document generated: {output_document_path}")
        with open(output_document_path, "rb") as file:
            st.download_button("Download Document", file, "Content_with_Images.docx")

##################################
    # Streamlit UI for merging video
    if st.button("Generate Video with Summary"):
        if new_content_data:
            merged_video_path = os.path.join(current_directory, "merged_video.mp4")
            merge_summary_with_video(temp_video_file_path, json.dumps(new_content_data), merged_video_path)
            st.video(merged_video_path)  # Display the merged video
            st.download_button("Download Merged Video", open(merged_video_path, "rb"), "merged_video.mp4")
        else:
            st.error("No summary available to overlay on the video.")
####################################

    # Clean up temporary files
    os.remove(temp_video_file_path)
